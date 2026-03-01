import os
import re
import sys
import tempfile
import datetime
import time
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class TrackChangesHelper:
    """
    Helper class to inject OpenXML <w:ins> and <w:del> tags for track changes support in python-docx.
    """
    _id_counter = 0

    @classmethod
    def _get_next_id(cls):
        """Generates a unique ID for track changes."""
        cls._id_counter += 1
        # Combine timestamp and counter to ensure uniqueness even in fast loops
        return str(int(time.time())) + str(cls._id_counter)

    @staticmethod
    def _get_iso_date():
        """Returns current date in ISO format without microseconds (Word compatible)."""
        return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    @staticmethod
    def create_ins_node(text: str, author: str = "AI_Auditor"):
        """Creates a <w:ins> node containing a run with text."""
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), TrackChangesHelper._get_next_id())
        ins.set(qn('w:author'), author)
        ins.set(qn('w:date'), TrackChangesHelper._get_iso_date())

        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        ins.append(r)
        return ins

    @staticmethod
    def create_del_node(text: str, author: str = "AI_Auditor"):
        """Creates a <w:del> node containing delText."""
        del_tag = OxmlElement('w:del')
        del_tag.set(qn('w:id'), TrackChangesHelper._get_next_id())
        del_tag.set(qn('w:author'), author)
        del_tag.set(qn('w:date'), TrackChangesHelper._get_iso_date())

        del_text = OxmlElement('w:delText')
        del_text.text = text
        del_tag.append(del_text)
        return del_tag

    @staticmethod
    def clear_paragraph_content_keep_properties(paragraph):
        """
        Clears all content from the paragraph but PRESERVES the pPr element (paragraph properties).
        This prevents loss of styling (alignment, indentation, etc.) when replacing text.
        """
        p = paragraph._p
        # We want to remove all children EXCEPT w:pPr
        # Iterate over a copy of children list
        for child in list(p):
            if child.tag != qn('w:pPr'):
                p.remove(child)

    @staticmethod
    def mark_paragraph_deleted(paragraph, author="AI_Auditor"):
        """Wraps the entire paragraph content in a delete marker."""
        text = paragraph.text
        if not text:
            return
        
        p = paragraph._p
        TrackChangesHelper.clear_paragraph_content_keep_properties(paragraph)
        
        # Add deletion node
        del_node = TrackChangesHelper.create_del_node(text, author)
        p.append(del_node)

    @staticmethod
    def mark_paragraph_replaced(paragraph, new_text, author="AI_Auditor"):
        """Marks old content as deleted and appends new content as inserted."""
        old_text = paragraph.text
        
        p = paragraph._p
        TrackChangesHelper.clear_paragraph_content_keep_properties(paragraph)
        
        # Add deletion node for old text
        if old_text:
            del_node = TrackChangesHelper.create_del_node(old_text, author)
            p.append(del_node)
        
        # Add insertion node for new text
        ins_node = TrackChangesHelper.create_ins_node(new_text, author)
        p.append(ins_node)

    @staticmethod
    def append_insertion(paragraph, text, author="AI_Auditor"):
        """Appends text to a paragraph as an insertion."""
        p = paragraph._p
        ins_node = TrackChangesHelper.create_ins_node(text, author)
        p.append(ins_node)


class WordProcessor:
    def __init__(self):
        self.my_country = "Philippines"

    def audit_and_fix(self, input_path, output_path):
        """执行红线审计并返回结果。全平台统一使用 XML 注入模式。"""
        input_abs_path = str(Path(input_path).resolve())
        output_abs_path = str(Path(output_path).resolve())

        print(f"[WordProcessor] 使用 XML 注入进行审计...")
        return self._audit_and_fix_xml(input_abs_path, output_abs_path)

    def _audit_and_fix_xml(self, input_abs_path: str, output_abs_path: str):
        """
        全平台通用方案：使用 python-docx + XML 注入实现修订模式。
        """
        audit_results = []
        try:
            doc = Document(input_abs_path)
            full_text = self._collect_doc_text(doc)

            audit_results.extend(self._check_sensitive_info_text(full_text))
            audit_results.extend(self._check_signatories_text(full_text))
            # 传递 doc 对象进行 XML 注入修改
            audit_results.extend(self._fix_payment_invoice_docx(doc, full_text))
            audit_results.extend(self._fix_dispute_clause_docx(doc))
            audit_results.extend(self._delete_penalty_docx(doc))

            Path(output_abs_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_abs_path)

            return audit_results
        except Exception as e:
            print(f"XML 审计过程出错: {e}")
            return [{"id": "err", "level": "error", "title": "审计引擎异常", "content": str(e), "anchor": ""}]

    @staticmethod
    def _collect_doc_text(doc: DocxDocument) -> str:
        paragraph_texts = [p.text for p in doc.paragraphs if p.text]
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    table_texts.append(row_text)
        return "\n".join(paragraph_texts + table_texts)

    @staticmethod
    def _check_sensitive_info_text(full_text: str):
        results = []
        phone_pattern = r"\+86\s?\d+"
        email_pattern = r"\b[A-Za-z0-9._%+-]+@(126|163)\.com\b"
        mark_counter = 0

        for pattern, msg in [
            (phone_pattern, "Please confirm: 中国电话 (+86)"),
            (email_pattern, "Please confirm: 126/163邮箱"),
        ]:
            for match in re.finditer(pattern, full_text):
                mark_id = f"mark_sensitive_{mark_counter}"
                mark_counter += 1
                results.append(
                    {
                        "id": mark_id,
                        "level": "warning",
                        "title": "敏感联系方式",
                        "content": f"识别到中国元素，建议确认：{msg}",
                        "anchor": match.group(),
                    }
                )
        return results

    @staticmethod
    def _check_signatories_text(full_text: str):
        results = []
        last_part = full_text[-500:]
        if "Position" not in last_part and "Title" not in last_part:
            results.append(
                {
                    "id": "sig_missing",
                    "level": "warning",
                    "title": "签字人职位缺失",
                    "content": "文末未检测到签字人职位抬头，请手动补全以符合合规要求。",
                    "anchor": "Signature",
                }
            )
        return results

    def _fix_payment_invoice_docx(self, doc: DocxDocument, full_text: str):
        results = []
        lower_text = full_text.lower()
        if "invoice" in lower_text and "prior to each payment" not in lower_text:
            for para in doc.paragraphs:
                if "payment" in para.text.lower():
                    # Use XML injection for track changes
                    TrackChangesHelper.append_insertion(
                        para, 
                        " Party B shall issue a valid and lawful invoice of the corresponding amount to Party A prior to each payment made by Party A."
                    )
                    results.append(
                        {
                            "id": "mark_payment_invoice",
                            "level": "error",
                            "title": "缺失发票逻辑",
                            "content": "检测到付款条款但缺失‘先票后款’，已自动修订补充。",
                            "anchor": "prior to each payment",
                        }
                    )
                    break
        return results

    def _fix_dispute_clause_docx(self, doc: DocxDocument):
        results = []
        standard_dispute = (
            f"This Agreement shall be governed by and construed in accordance with the laws of {self.my_country}. "
            "Any dispute shall be submitted to the exclusive jurisdiction of the competent courts of Party A."
        )

        for para in doc.paragraphs:
            if "DISPUTE RESOLUTION" in para.text:
                new_content = "DISPUTE RESOLUTION\n" + standard_dispute
                # Use XML injection to show replacement
                TrackChangesHelper.mark_paragraph_replaced(para, new_content)
                
                results.append(
                    {
                        "id": "mark_dispute_resolution",
                        "level": "error",
                        "title": "争议解决修订",
                        "content": f"已将管辖权自动替换为我方所在地 ({self.my_country})。",
                        "anchor": "DISPUTE RESOLUTION",
                    }
                )
                break
        return results

    @staticmethod
    def _delete_penalty_docx(doc: DocxDocument):
        results = []
        patterns = ["late payment penalty", "penalty interest"]
        penalty_counter = 0

        for para in doc.paragraphs:
            original_text = para.text
            updated_text = original_text
            for pattern in patterns:
                if pattern in updated_text.lower():
                    mark_id = f"mark_penalty_{penalty_counter}"
                    penalty_counter += 1
                    results.append(
                        {
                            "id": mark_id,
                            "level": "error",
                            "title": "违约金删除",
                            "content": f"识别到罚息表述 '{pattern}'，已自动删除。",
                            "anchor": original_text[:30] if original_text else pattern,
                        }
                    )
                    updated_text = re.sub(pattern, "", updated_text, flags=re.IGNORECASE)
            
            if updated_text != original_text:
                # Use XML injection to show deletion
                # Logic: We replace the WHOLE paragraph with the cleaned text.
                # Ideally we should only delete the specific span, but replacing whole para is safer for now.
                TrackChangesHelper.mark_paragraph_replaced(para, updated_text.strip())

        return results
