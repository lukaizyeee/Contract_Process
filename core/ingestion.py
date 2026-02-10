import re
from dataclasses import dataclass
from typing import List, Optional
from docx import Document
from pathlib import Path

@dataclass
class Chunk:
    text: str
    original_index: int
    source_type: str  # 'paragraph' or 'table'
    metadata: Optional[dict] = None

class DocProcessor:
    def __init__(self, max_chars: int = 400, window_size: int = 4, overlap: int = 1):
        self.max_chars = max_chars
        self.window_size = window_size  # Number of sentences per chunk
        self.overlap = overlap  # Number of overlapping sentences

    def process(self, file_path: str) -> List[Chunk]:
        """
        Process a .docx file and return a list of Chunks.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() == '.doc':
             raise ValueError("Old .doc format is not supported. Please save as .docx.")

        doc = Document(file_path)
        chunks = []
        
        # Process paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = self._clean_text(para.text)
            if not text:
                continue
                
            if len(text) > self.max_chars:
                chunks.extend(self._sliding_window(text, i, 'paragraph'))
            else:
                chunks.append(Chunk(text=text, original_index=i, source_type='paragraph'))
                
        # Process tables
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                # Combine cell texts with separator
                cell_texts = [self._clean_text(cell.text) for cell in row.cells]
                # Remove empty cells from check? Or keep them? 
                # Keeping them maintains structure.
                row_text = " | ".join(cell_texts)
                
                # Check if row is essentially empty
                if not self._clean_text(row_text).replace('|', '').strip():
                    continue
                    
                # For tables, we treat each row as a unit usually, 
                # but if a cell is very long, we might want to split it?
                # For now, treat row as atomic as per plan implication (simple cleaning).
                # Plan says "读取...段落...和表格". 
                # We use source_type to store table index and row index.
                # We can store 'table_{i}' as source_type and j as extra info, 
                # or just 'table' and metadata={'table_index': i, 'row_index': j}
                
                chunks.append(Chunk(
                    text=row_text, 
                    original_index=i, 
                    source_type='table',
                    metadata={'row_index': j}
                ))

        return chunks

    def _clean_text(self, text: str) -> str:
        """
        Remove empty lines, header/footer artifacts (simple whitespace cleaning).
        """
        # Normalize whitespace
        return re.sub(r'\s+', ' ', text).strip()

    def _sliding_window(self, text: str, index: int, source_type: str) -> List[Chunk]:
        """
        Split long text into overlapping chunks of sentences.
        """
        # Split into sentences (simple regex)
        # Look for [.!?] followed by space or end of line
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Remove empty strings
        sentences = [s for s in sentences if s.strip()]
        
        if not sentences:
            return []

        if len(sentences) <= self.window_size:
            return [Chunk(text=text, original_index=index, source_type=source_type)]
            
        chunks = []
        stride = max(1, self.window_size - self.overlap)
        
        for i in range(0, len(sentences), stride):
            window = sentences[i : i + self.window_size]
            chunk_text = " ".join(window)
            
            # Avoid very small chunks at the end if possible, or just accept them
            if chunk_text:
                chunks.append(Chunk(text=chunk_text, original_index=index, source_type=source_type))
            
            if i + self.window_size >= len(sentences):
                break
                
        return chunks

if __name__ == "__main__":
    # Test stub
    pass
