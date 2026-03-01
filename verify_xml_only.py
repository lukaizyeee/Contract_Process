import os
import sys
from core.preview_generator import DocxPreviewGenerator
from core.word_processor import WordProcessor
from docx import Document

def create_dummy_docx(filename):
    doc = Document()
    doc.add_paragraph("This is a test document for XML verification.")
    doc.add_paragraph("It contains some text to simulate a contract.")
    doc.save(filename)
    return filename

def test_preview_generator(filename):
    print("\n--- Testing Preview Generator ---")
    generator = DocxPreviewGenerator()
    html_content = generator.generate_html(filename)
    if "<html>" in html_content and "This is a test document" in html_content:
        print("SUCCESS: HTML generated successfully via XML parsing.")
    else:
        print("FAILURE: HTML generation failed or content missing.")
        print(f"Content snippet: {html_content[:100]}...")

def test_word_processor(filename):
    print("\n--- Testing Word Processor ---")
    processor = WordProcessor()
    output_filename = filename.replace(".docx", "_revised.docx")
    
    # Run audit (even if no changes found, it should run without error)
    results = processor.audit_and_fix(filename, output_filename)
    print(f"Audit finished. Results found: {len(results)}")
    
    if os.path.exists(output_filename):
        print(f"SUCCESS: Output file created at {output_filename}")
        # Verify it's a valid docx
        try:
            doc = Document(output_filename)
            print("SUCCESS: Output file is a valid DOCX.")
        except Exception as e:
            print(f"FAILURE: Output file is corrupt: {e}")
    else:
        print("FAILURE: Output file not created.")

if __name__ == "__main__":
    test_file = "verify_xml.docx"
    create_dummy_docx(test_file)
    
    try:
        test_preview_generator(test_file)
        test_word_processor(test_file)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
        if os.path.exists(test_file.replace(".docx", "_revised.docx")):
            os.remove(test_file.replace(".docx", "_revised.docx"))
