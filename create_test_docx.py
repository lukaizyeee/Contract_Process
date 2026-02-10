from docx import Document

doc = Document()
doc.add_heading('Service Agreement', 0)

doc.add_paragraph('This Agreement is made on 2024-01-01 between Party A and Party B.')

doc.add_heading('1. Term', level=1)
doc.add_paragraph('This agreement shall commence on the Effective Date and shall continue for a period of one (1) year unless terminated earlier in accordance with the provisions of this Agreement.')

doc.add_heading('2. Termination', level=1)
doc.add_paragraph('Either party may terminate this Agreement upon written notice if the other party breaches any material term of this Agreement and fails to cure such breach within thirty (30) days of receipt of such notice.')
doc.add_paragraph('In the event of termination, Party B shall return all confidential information to Party A immediately.')

doc.add_heading('3. Confidentiality', level=1)
doc.add_paragraph('Each party agrees to keep confidential all non-public information received from the other party.')

doc.add_heading('4. Payment', level=1)
table = doc.add_table(rows=3, cols=2)
table.rows[0].cells[0].text = 'Item'
table.rows[0].cells[1].text = 'Price'
table.rows[1].cells[0].text = 'Consulting Service'
table.rows[1].cells[1].text = '$100/hr'
table.rows[2].cells[0].text = 'Software License'
table.rows[2].cells[1].text = '$5000/year'

doc.save('test_contract.docx')
print("Created test_contract.docx")
